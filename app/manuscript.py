"""EL YAZMASI DIŞA AKTARIMI: romanı okunur biçimde indirmek.

Mevcut JSON yedeği bir VERİ yedeğidir - geri yüklemek için, okumak için
değil. Romanı basmak, birine göndermek, bir editöre vermek ya da başka
bir programda açmak için okunur bir çıktı gerekiyordu (README'nin bilinen
eksikler listesinde ilk sıradaydı).

Fihrist hiyerarşisi korunur: KISIM > bölüm > alt başlık. Paragraflar sıra
numarasıyla değil düz metin olarak, okunacak gibi dizilir.
"""
from io import BytesIO

from . import models
from .outline import build_hierarchy

# Word'de hangi hiyerarşi seviyesi hangi başlık düzeyi olsun.
SEVIYE_BASLIK = {1: 1, 2: 2, 3: 3, 4: 4}


def _bolumler(db, novel_id: int):
    return (
        db.query(models.Chapter)
        .filter(models.Chapter.novel_id == novel_id)
        .order_by(models.Chapter.number)
        .all()
    )


def _paragraflar(chapter):
    return sorted(
        [p for p in chapter.paragraphs if (p.text or "").strip()],
        key=lambda p: p.number,
    )


def export_markdown(db, novel) -> str:
    """Markdown: her yerde açılır, sürüm takibine girer, düz metindir."""
    satirlar = [f"# {novel.name}", ""]
    for item in build_hierarchy(_bolumler(db, novel.id)):
        ch = item["chapter"]
        seviye = min(item["level"], 4) + 1  # roman adı zaten h1
        baslik = (ch.title or "").strip() or f"Bölüm {ch.number}"
        satirlar.append(f"{'#' * seviye} {baslik}")
        satirlar.append("")
        for p in _paragraflar(ch):
            satirlar.append(p.text.strip())
            satirlar.append("")
    return "\n".join(satirlar).rstrip() + "\n"


def export_docx(db, novel) -> bytes:
    """Word: yazdırmak, paylaşmak, üzerine not almak için.

    Başlıklar GERÇEK Word başlığı olduğu için belge içi gezinme ve
    içindekiler tablosu çalışır.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    belge = Document()
    baslik = belge.add_heading(novel.name, level=0)
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Roman gövdesi için okunaklı bir taban: 12 punto, satır arası geniş.
    normal = belge.styles["Normal"]
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.3

    for item in build_hierarchy(_bolumler(db, novel.id)):
        ch = item["chapter"]
        metin = (ch.title or "").strip() or f"Bölüm {ch.number}"
        belge.add_heading(metin, level=SEVIYE_BASLIK.get(item["level"], 4))
        for p in _paragraflar(ch):
            belge.add_paragraph(p.text.strip())

    tampon = BytesIO()
    belge.save(tampon)
    return tampon.getvalue()


def export_txt(db, novel) -> str:
    """Düz metin: başlık işaretleri olmadan, sadece roman."""
    parcalar = [novel.name, ""]
    for item in build_hierarchy(_bolumler(db, novel.id)):
        ch = item["chapter"]
        baslik = (ch.title or "").strip() or f"Bölüm {ch.number}"
        parcalar.append(baslik.upper() if item["level"] == 1 else baslik)
        parcalar.append("")
        for p in _paragraflar(ch):
            parcalar.append(p.text.strip())
            parcalar.append("")
    return "\n".join(parcalar).rstrip() + "\n"


def istatistik(db, novel_id: int) -> dict:
    """İndirmeden önce ne kadar metin olduğunu göstermek için."""
    bolumler = _bolumler(db, novel_id)
    paragraf = kelime = karakter = 0
    yazili_bolum = 0
    for ch in bolumler:
        ps = _paragraflar(ch)
        if ps:
            yazili_bolum += 1
        paragraf += len(ps)
        for p in ps:
            karakter += len(p.text)
            kelime += len(p.text.split())
    return {
        "girdi": len(bolumler),
        "yazili_bolum": yazili_bolum,
        "paragraf": paragraf,
        "kelime": kelime,
        "karakter": karakter,
    }
