"""MATRİS DENETİM PROMPTU: tamamlanmış bir planı dışarıdan denetletmek için.

İki iş yapar ve ayrımı önemlidir:

1. DETERMİNİSTİK KONTROLLER (burada, AI'sız, ücretsiz): sayılabilir,
   kesin cevabı olan sorular - boş hücre var mı, dolu ama hiçbir bölüme
   bağlanmamış plan var mı (o plan AI'ya HİÇ gitmez), var olmayan bir MP
   koduna referans verilmiş mi, aynı bölüme iki hücre birden bağlanmış mı.
   Bunlar bir modele sorulmaz; sorulursa hem para yakar hem yanılır.

2. PROMPT (dışarıya): geri kalanı anlam gerektirir - talimat kasasıyla
   beat'in çelişip çelişmediği, turların gerçekten paralel olup olmadığı,
   damganın düşüp düşmediği. Bunlar için plan okunabilir metne dökülüp
   sorularla birlikte paketlenir; kullanıcı istediği modele yapıştırır.

Yani buradaki çıktı bir "AI özelliği" değil: mevcut verinin denetlenebilir
bir dökümü. Hiçbir Qwen çağrısı yapılmaz.
"""
from . import models
from . import plan_schema


def insan_metni(content: str) -> str:
    """Plan metnini İNSAN için temizler: AI'ya yönelik yönergeleri atar.

    content, modele giden metindir - içinde SINIRLAR bloğu, "sahne BU AN'da
    geçer" uyarısı, hedef uzunluğun tarifi gibi MODELE söylenen şeyler var.
    Basılı bir planda bunlar gürültüdür; yazar kendi planını okumak ister,
    modele verilen talimatı değil.
    """
    if not content:
        return ""
    satirlar = []
    for satir in content.split("\n"):
        # SINIRLAR bloğu ve altındaki maddeler tamamen çıkar.
        if satir.startswith("SINIRLAR"):
            break
        # Uzun yönerge kuyruklarını kes, bilgiyi bırak.
        if satir.startswith("ZAMAN:") and " — sahne BU AN" in satir:
            satir = satir.split(" — sahne BU AN")[0]
        elif satir.startswith("ODAK:") and " — betimleme" in satir:
            satir = satir.split(" — betimleme")[0]
        elif satir.startswith("HEDEF UZUNLUK:"):
            satir = satir.split(".")[0]
        satirlar.append(satir)
    return "\n".join(satirlar).rstrip()


def _kolon_basligi(col) -> str:
    tur = plan_schema.normalize_meta(getattr(col, "tur_data", None), plan_schema.TUR_ALANLARI)
    dolu = [(etiket, tur[key]) for key, etiket, _ in plan_schema.TUR_ALANLARI if tur[key]]
    if not dolu:
        return f"### {col.label}\n(TUR MİRASI tanımlanmamış)"
    return f"### {col.label}\nTUR MİRASI: " + " · ".join(f"{e}: {v}" for e, v in dolu)


def yapisal_bulgular(db, matrix, novel_id: int, columns, rows, cells) -> list[str]:
    """Sayılabilir, kesin kusurlar. Modele sorulmaz - burada bulunur."""
    bulgular = []
    hucre_map = {(c.column_id, c.row_id): c for c in cells}
    kodlar = {c.code for c in cells if c.code}

    # Roman genelindeki tüm MP kodları: referanslar buraya bakmalı.
    tum_kodlar = {
        kod for (kod,) in db.query(models.MatrixCell.code)
        .join(models.PlanMatrix, models.MatrixCell.matrix_id == models.PlanMatrix.id)
        .filter(models.PlanMatrix.novel_id == novel_id).all() if kod
    }

    bos = [(c.label, r.label) for c in columns for r in rows
           if not (hucre_map.get((c.id, r.id)) and (hucre_map[(c.id, r.id)].content or "").strip())]
    if bos:
        ornek = ", ".join(f"{k}×{s}" for k, s in bos[:6])
        bulgular.append(f"BOŞ HÜCRE: {len(bos)} kesişim doldurulmamış ({ornek}{'…' if len(bos) > 6 else ''})")

    # En ağır kusur: plan yazılmış ama hiçbir bölüme bağlanmamış. O plan
    # yazım sırasında AI'ya HİÇ gitmez - emek görünmez yere akar.
    bagsiz = [c for c in cells if (c.content or "").strip() and not c.chapter_id]
    if bagsiz:
        bulgular.append(
            f"BAĞSIZ PLAN: {len(bagsiz)} dolu hücre hiçbir bölüme bağlı değil "
            f"({', '.join(x.code or '?' for x in bagsiz[:8])}) - bu planlar yazım sırasında AI'ya GİTMEZ")

    # Aynı bölüme iki hücre: hangisinin planı geçerli belirsiz.
    bolum_sayaci = {}
    for c in cells:
        if c.chapter_id:
            bolum_sayaci.setdefault(c.chapter_id, []).append(c.code or "?")
    cakisan = {b: k for b, k in bolum_sayaci.items() if len(k) > 1}
    if cakisan:
        bulgular.append("ÇİFT BAĞ: aynı bölüme birden çok hücre bağlı - "
                        + "; ".join("+".join(k) for k in cakisan.values()))

    # Var olmayan koda referans / kendine referans.
    for c in cells:
        veri = plan_schema.normalize_cell(c.data)
        for b in veri["baglantilar"]:
            if b["kod"] == c.code:
                bulgular.append(f"KENDİNE REFERANS: {c.code} kendi kodunu gösteriyor")
            elif b["kod"] not in tum_kodlar:
                bulgular.append(f"KAYIP REFERANS: {c.code or '?'} → {b['kod']} diye bir hücre yok")

    # Damgası tanımlanmamış tur: damga kilidi o kolonda hiç çalışmaz.
    for col in columns:
        tur = plan_schema.normalize_meta(getattr(col, "tur_data", None), plan_schema.TUR_ALANLARI)
        if not tur.get("damga"):
            bulgular.append(f"DAMGASIZ TUR: \"{col.label}\" için damga tanımlanmamış - "
                            f"o turun hiçbir hücresinde damga kontrolü çalışmaz")

    # Paralellik: turlar yapıca aynı doğrultuda ilerlemek zorunda.
    bulgular.extend(paralellik_bulgulari(columns, rows, cells))

    return bulgular


def build_audit_prompt(db, matrix, novel_id: int, column_id: int | None = None) -> tuple[str, dict]:
    """Denetim promptunu ve özet sayıları üretir. column_id verilirse tek
    tur denetlenir - 8×7'lik bir matrisin tamamı tek isteğe sığmaz ve
    kullanıcı zaten tur tur çalışır."""
    columns = [c for c in matrix.columns if column_id is None or c.id == column_id]
    rows = list(matrix.rows)
    ilgili_kolonlar = {c.id for c in columns}
    cells = [c for c in matrix.cells if c.column_id in ilgili_kolonlar]

    hucre_map = {(c.column_id, c.row_id): c for c in cells}
    cols_by_id = {c.id: c for c in matrix.columns}

    bulgular = yapisal_bulgular(db, matrix, novel_id, columns, rows, cells)

    # Hücre bazlı uyarılar (mevcut yapı kilidi denetimi)
    hucre_uyarilari = []
    for c in cells:
        uyarilar = plan_schema.cell_warnings(
            c.data, getattr(cols_by_id.get(c.column_id), "tur_data", None),
            paralel=len(matrix.columns) > 1)
        if uyarilar:
            hucre_uyarilari.append(f"- {c.code or '?'}: " + " · ".join(uyarilar))

    # --- Prompt metni ---
    p = []
    p.append("# PLAN DENETİMİ — YAPI KİLİDİ v1.2")
    p.append("")
    p.append(f"Aşağıda \"{matrix.name}\" planının "
             + (f"tek turu ({columns[0].label})" if column_id else f"{len(columns)} turu")
             + f" ve {len(rows)} aşaması var. Bu planı YAZMA - DENETLE.")
    p.append("")
    p.append("Yapı şu: MİRAS alanları hücrede değil, turun/parçanın kaydında durur ve "
             "her sahneye canlı eklenir. Hücre ise sabit bir sahne şemasıdır: "
             "OLAY · ZAMAN · MEKAN · DUYGU · ODAK üstte, GİRİŞ/GELİŞME/SONUÇ yayı altta, "
             "MP bağlantıları en sonda. GİRİŞ/GELİŞME/SONUÇ birer BEAT'tir - tek bir an, "
             "olay dizisi değil.")
    p.append("")

    p.append("## SİSTEMİN KENDİ BULDUKLARI")
    p.append("")
    p.append("Bunlar zaten tespit edildi; tekrar saymana gerek yok. Sana bunların "
             "**sonucunu** soracağım.")
    p.append("")
    if bulgular:
        p.extend(f"- {b}" for b in bulgular)
    else:
        p.append("- Yapısal kusur bulunamadı.")
    p.append("")
    if hucre_uyarilari:
        p.append("Hücre bazlı eksikler:")
        p.append("")
        p.extend(hucre_uyarilari)
    else:
        p.append("Hücre bazlı eksik yok.")
    p.append("")

    p.append("## PLAN")
    p.append("")
    for col in columns:
        p.append(_kolon_basligi(col))
        p.append("")
        for row in rows:
            cell = hucre_map.get((col.id, row.id))
            metin = (cell.content or "").strip() if cell else ""
            parca = plan_schema.normalize_meta(getattr(row, "parca_data", None), plan_schema.PARCA_ALANLARI)
            parca_ek = " · ".join(f"{e}: {parca[k]}" for k, e, _ in plan_schema.PARCA_ALANLARI if parca[k])
            baslik = f"**{row.label}**" + (f" ({parca_ek})" if parca_ek else "")
            kod = f" [{cell.code}]" if cell and cell.code else ""
            # Bağ durumu denetim için kritik: bağsız plan AI'ya hiç gitmez.
            bolum = ""
            if cell and cell.chapter_id:
                ch = db.query(models.Chapter).filter(models.Chapter.id == cell.chapter_id).first()
                bolum = f" → Bölüm {ch.number}" if ch else ""
            elif metin:
                bolum = " → (BAĞSIZ - bu plan AI'ya gitmez)"
            p.append(baslik + kod + bolum)
            if (row.instructions or "").strip():
                p.append(f"YAZIM KISITLARI: {row.instructions.strip()}")
            p.append(metin if metin else "(boş)")
            p.append("")

    p.append("## SORULAR")
    p.append("")
    p.append("**1. Çelişki.** Herhangi bir hücrenin beat'i, kendi satırının YAZIM "
             "KISITLARI ile aynı anda uygulanabilir mi? Uygulanamayan varsa hangi iki "
             "satırın çakıştığını göster. (Bu, sistemin bulamadığı türden bir kusurdur - "
             "anlam gerektirir.)")
    p.append("")
    p.append("**2b. Uzunluk.** Her hücrenin HEDEF UZUNLUK ölçüsü, o hücrenin "
             "beat'lerinin taşıyabileceği yükle uyuşuyor mu? Üç beat'i olan bir "
             "sahneye \"özet\" denmişse ya da tek jestlik bir ana \"uzun metin\" "
             "denmişse göster.")
    p.append("")
    p.append("**2. Beat mi, dizi mi?** Hangi GİRİŞ/GELİŞME/SONUÇ alanları bir AN değil "
             "de olay dizisi olmuş? Bunları yazmaya kalkan bir modele kuracak yer kalıyor mu?")
    p.append("")
    p.append("**3. Damga.** Her turun damga kelimesi, o turun hücrelerinin SONUÇ "
             "beat'lerinde gerçekten DÜŞÜYOR mu, yoksa cümleye iliştirilmiş mi? "
             "Kelimenin geçmesi yetmez - dramatik olarak düşmesi gerekir.")
    p.append("")
    p.append("**4. Paralellik.** Aynı aşamanın farklı turlardaki hâlleri birbirinin "
             "tekrarı mı, yoksa yükselen bir çizgi mi? Tekrar olan çiftleri göster.")
    p.append("")
    p.append("**5. Bağlantılar.** MP bağlantıları yapılacak işi söylüyor mu, yoksa "
             "sadece referans mı veriyor? Referans kalmış olanları listele.")
    p.append("")
    p.append("**6. Yukarıdaki 'SİSTEMİN KENDİ BULDUKLARI' listesinden hangisi gerçekten "
             "ciddi, hangisi gürültü?** Sıraya koy - önce hangisini düzelteyim?")
    p.append("")
    p.append("**7.** Bu planla sahneleri yazabilir misin, yoksa doldurur musun? "
             "Aradaki farkı bu plandan bir örnekle göster.")

    ozet = {
        "kolon_sayisi": len(columns),
        "satir_sayisi": len(rows),
        "dolu_hucre": sum(1 for c in cells if (c.content or "").strip()),
        "yapisal_bulgu": len(bulgular),
        "uyarili_hucre": len(hucre_uyarilari),
    }
    return "\n".join(p), ozet


# --- TOPLU DIŞA AKTARIM ----------------------------------------------------

def _hucre_sozlugu(db, cell, col, row) -> dict:
    """Bir hücrenin tam kaydı. content YENİDEN ÜRETİLMEZ, olduğu gibi
    taşınır: eski serbest metinli hücrelerin içeriği de kaybolmasın."""
    ch = None
    if cell.chapter_id:
        ch = db.query(models.Chapter).filter(models.Chapter.id == cell.chapter_id).first()
    return {
        "kod": cell.code,
        "kolon": col.label if col else None,
        "satir": row.label if row else None,
        "bolum": ch.number if ch else None,
        "bolum_basligi": (ch.title if ch else None),
        "veri": plan_schema.normalize_cell(cell.data),
        "metin": cell.content or "",
        "uyarilar": plan_schema.cell_warnings(
            cell.data, getattr(col, "tur_data", None) if col else None),
    }


def export_json(db, matrisler, novel_id: int) -> dict:
    """Makine okunur tam döküm - yeniden içeri alınabilecek şekilde."""
    from datetime import datetime, timezone
    cikti = {
        "surum": "yapi-kilidi-1.2",
        "alindi": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "matrisler": [],
    }
    for m in matrisler:
        cols_by_id = {c.id: c for c in m.columns}
        rows_by_id = {r.id: r for r in m.rows}
        cikti["matrisler"].append({
            "ad": m.name,
            "turlar": [{
                "etiket": c.label,
                "tur_mirasi": plan_schema.normalize_meta(
                    getattr(c, "tur_data", None), plan_schema.TUR_ALANLARI),
            } for c in m.columns],
            "asamalar": [{
                "etiket": r.label,
                "tur": r.kind or "main",
                "yazim_kisitlari": r.instructions or "",
                "parca_mirasi": plan_schema.normalize_meta(
                    getattr(r, "parca_data", None), plan_schema.PARCA_ALANLARI),
            } for r in m.rows],
            "hucreler": [
                _hucre_sozlugu(db, c, cols_by_id.get(c.column_id), rows_by_id.get(c.row_id))
                for c in m.cells
            ],
        })
    return cikti


def export_markdown(db, matrisler, novel_id: int) -> str:
    """İnsan okunur döküm - okumak, yazdırmak, başka bir modele vermek için."""
    p = ["# PLAN MATRİSLERİ", ""]
    for m in matrisler:
        cols_by_id = {c.id: c for c in m.columns}
        hucre_map = {(c.column_id, c.row_id): c for c in m.cells}
        p.append(f"## {m.name}")
        p.append("")
        for col in m.columns:
            p.append(_kolon_basligi(col))
            p.append("")
            for row in m.rows:
                cell = hucre_map.get((col.id, row.id))
                metin = insan_metni((cell.content or "").strip() if cell else "")
                parca = plan_schema.normalize_meta(
                    getattr(row, "parca_data", None), plan_schema.PARCA_ALANLARI)
                ek = " · ".join(f"{e}: {parca[k]}"
                                for k, e, _ in plan_schema.PARCA_ALANLARI if parca[k])
                kod = f" [{cell.code}]" if cell and cell.code else ""
                bolum = ""
                if cell and cell.chapter_id:
                    ch = db.query(models.Chapter).filter(
                        models.Chapter.id == cell.chapter_id).first()
                    bolum = f" → Bölüm {ch.number}" if ch else ""
                p.append(f"**{row.label}**" + (f" ({ek})" if ek else "") + kod + bolum)
                if (row.instructions or "").strip():
                    p.append(f"YAZIM KISITLARI: {row.instructions.strip()}")
                p.append(metin if metin else "(boş)")
                p.append("")
    return "\n".join(p)


# --- PARALELLİK DENETİMİ ---------------------------------------------------

def paralellik_bulgulari(columns, rows, cells) -> list[str]:
    """Sütunlar (turlar) yapıca paralel ilerlemek zorunda: aynı satır her
    turda aynı yapısal konumdur. Satırlar ortak olduğu için sahne SAYISI
    zaten garanti - ızgara bunu fiziksel olarak dayatır.

    Garanti EDİLMEYEN iki şey var ve ikisi de sessiz:
      1. Hücre boş bırakılabilir. Tur 1'de dolu, Tur 2'de boş bir satır,
         yazımda "biri 8 sahne biri 5 sahne" olarak geri döner.
      2. Aynı satırda farklı HEDEF UZUNLUK seçilebilir. Paralel olması
         gereken iki sahne, biri 2 paragraf biri 10 paragraf çıkar.

    BAŞLANMAMIŞ sütunlar denetim dışıdır: hiç dokunulmamış Tur 5'i
    "eksik" diye raporlamak, gerçek boşlukları gürültüye gömer.
    """
    bulgular = []
    dolu = {(c.column_id, c.row_id): c for c in cells if (c.content or "").strip()}
    baslanmis = [c for c in columns if any(k[0] == c.id for k in dolu)]
    if len(baslanmis) < 2:
        return bulgular  # karşılaştıracak ikinci tur yok

    # 1. Delik: başlanmış bir turda boş kalan satırlar
    for col in baslanmis:
        eksik = [r.label for r in rows if (col.id, r.id) not in dolu]
        if eksik:
            ornek = ", ".join(eksik[:5])
            bulgular.append(
                f"PARALELLİK DELİĞİ: \"{col.label}\" {len(eksik)} satırda boş "
                f"({ornek}{'…' if len(eksik) > 5 else ''}) - diğer turlarda dolu olan "
                f"bu konumlar yazımda eksik sahne olarak çıkar")

    # 2. Uzunluk uyuşmazlığı: aynı satır, farklı ölçü
    for row in rows:
        olculer = {}
        for col in baslanmis:
            h = dolu.get((col.id, row.id))
            if h:
                seviye = plan_schema.normalize_cell(h.data)["uzunluk"]
                olculer.setdefault(seviye, []).append(col.label)
        if len(olculer) > 1:
            detay = " / ".join(f"{s}: {', '.join(k)}" for s, k in olculer.items())
            bulgular.append(
                f"UZUNLUK UYUŞMAZLIĞI: \"{row.label}\" satırı turlar arasında farklı "
                f"ölçüde ({detay}) - paralel sahneler farklı boyda çıkar")

    return bulgular


def export_docx(db, matrisler, novel_id: int) -> bytes:
    """Word belgesi: yazdırmak, paylaşmak, üzerine elle not almak için.

    Markdown dökümüyle aynı bilgiyi taşır ama BAŞLIK DÜZEYLERİ gerçek Word
    başlıkları olduğu için belge içi gezinme ve içindekiler tablosu çalışır.
    Boş hücreler de yazılır - hangi kesişimin doldurulmadığı basılı planda
    da görünmeli.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    belge = Document()

    baslik = belge.add_heading("Plan Matrisleri", level=0)
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for m in matrisler:
        belge.add_heading(m.name, level=1)
        cols_by_id = {c.id: c for c in m.columns}
        hucre_map = {(c.column_id, c.row_id): c for c in m.cells}

        for col in m.columns:
            belge.add_heading(col.label, level=2)

            tur = plan_schema.normalize_meta(
                getattr(col, "tur_data", None), plan_schema.TUR_ALANLARI)
            dolu = [(e, tur[k]) for k, e, _ in plan_schema.TUR_ALANLARI if tur[k]]
            if dolu:
                p = belge.add_paragraph()
                r = p.add_run("TUR MİRASI: ")
                r.bold = True
                p.add_run(" · ".join(f"{e}: {v}" for e, v in dolu))
                p.runs[-1].font.size = Pt(9)

            for row in m.rows:
                cell = hucre_map.get((col.id, row.id))
                metin = insan_metni((cell.content or "").strip() if cell else "")

                parca = plan_schema.normalize_meta(
                    getattr(row, "parca_data", None), plan_schema.PARCA_ALANLARI)
                ek = " · ".join(f"{e}: {parca[k]}"
                                for k, e, _ in plan_schema.PARCA_ALANLARI if parca[k])
                girinti = "    " if (row.kind or "main") == "sub" else ""
                satir_basligi = f"{girinti}{row.label}"
                if cell and cell.code:
                    satir_basligi += f"  [{cell.code}]"
                if cell and cell.chapter_id:
                    ch = db.query(models.Chapter).filter(
                        models.Chapter.id == cell.chapter_id).first()
                    if ch:
                        satir_basligi += f"  → Bölüm {ch.number}"
                if ek:
                    satir_basligi += f"  ({ek})"
                belge.add_heading(satir_basligi, level=3)

                kurallar = (row.instructions or "").strip()
                if kurallar:
                    p = belge.add_paragraph()
                    r = p.add_run("YAZIM KISITLARI: ")
                    r.bold = True
                    p.add_run(kurallar)
                    for run in p.runs:
                        run.font.size = Pt(9)

                if not metin:
                    p = belge.add_paragraph("(boş)")
                    p.runs[0].italic = True
                    continue

                # Plan satırları: "ETİKET: içerik" - etiket kalın olsun ki
                # basılı sayfada göz satırları ayırt edebilsin.
                for satir in metin.split("\n"):
                    p = belge.add_paragraph()
                    if ":" in satir and satir.split(":", 1)[0].isupper():
                        etiket, icerik = satir.split(":", 1)
                        r = p.add_run(f"{etiket}:")
                        r.bold = True
                        p.add_run(icerik)
                    else:
                        p.add_run(satir)
                    for run in p.runs:
                        run.font.size = Pt(10)

    tampon = BytesIO()
    belge.save(tampon)
    return tampon.getvalue()
