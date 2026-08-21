

def test_paragraph_can_be_added_to_heading_that_already_has_text(client, headers):
    """REGRESYON: metni OLAN bir başlık girdisine paragraf eklenebilmeli.
    Kural "boş başlığa yanlışlıkla yazma" içindi; ama zaten metin taşıyan
    başlıklarda (içe aktarma, sonradan tür değişimi) PARAGRAF BÖLME ve
    taşıma bu engele takılıyordu."""
    ch = client.post("/chapters/", json={"number": 1, "kind": "chapter", "title": "B"}, headers=headers).json()
    client.put(f"/chapters/{ch['id']}/paragraphs/1", json={"number": 1, "text": "Uzun tek paragraf."}, headers=headers)
    client.put(f"/chapters/{ch['id']}", json={"kind": "subtitle"}, headers=headers)   # başlığa çevrildi

    # Bölme senaryosu: ikinci paragraf eklenebilmeli
    r = client.put(f"/chapters/{ch['id']}/paragraphs/2", json={"number": 2, "text": "İkinci parça."}, headers=headers)
    assert r.status_code == 200, r.text
    guncel = client.get(f"/chapters/{ch['id']}", headers=headers).json()
    assert len(guncel["paragraphs"]) == 2

    # BOŞ başlık hâlâ korunur - yanlışlıkla metin yazılmasın
    bos = client.post("/chapters/", json={"number": 5, "kind": "part", "title": "Ayraç"}, headers=headers).json()
    r2 = client.put(f"/chapters/{bos['id']}/paragraphs/1", json={"number": 1, "text": "Metin"}, headers=headers)
    assert r2.status_code == 400
    assert "henüz metni yok" in r2.json()["detail"]


def test_empty_paragraph_text_rejected(client, headers):
    """VERİ KAYBI KORUMASI: AI boş yanıt döndürdüğünde ya da arayüz
    hatasında paragraf içeriği sessizce siliniyordu. Geçmişten geri
    alınabiliyor ama kullanıcı fark etmeyebilir."""
    ch = client.post("/chapters/", json={"number": 1, "kind": "chapter", "title": "B"}, headers=headers).json()
    client.put(f"/chapters/{ch['id']}/paragraphs/1", json={"number": 1, "text": "Gerçek metin."}, headers=headers)

    for bos in ("", "   ", "\n\n"):
        r = client.put(f"/chapters/{ch['id']}/paragraphs/1", json={"number": 1, "text": bos}, headers=headers)
        assert r.status_code == 400, f"boş metin ({bos!r}) reddedilmeli"

    # Metin korundu
    guncel = client.get(f"/chapters/{ch['id']}", headers=headers).json()
    assert guncel["paragraphs"][0]["text"] == "Gerçek metin."


def test_paragraph_numbers_are_compacted_after_delete(client, headers):
    """Silinen numara BOŞLUK olarak kalıyordu: bütün paragrafları silsen
    bile yeni paragraf 9'dan başlayabiliyordu. Numara konumdur, kimlik değil."""
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    for i in range(1, 5):
        client.put(f"/chapters/{ch['id']}/paragraphs/{i}",
                   json={"number": i, "text": f"Paragraf {i}"}, headers=headers)

    # Ortadan sil -> kalanlar 1,2,3 olarak sıkışmalı
    assert client.delete(f"/chapters/{ch['id']}/paragraphs/2", headers=headers).status_code == 204
    kalan = client.get(f"/chapters/{ch['id']}", headers=headers).json()["paragraphs"]
    assert [p["number"] for p in kalan] == [1, 2, 3], "silme sonrası boşluk kaldı"
    assert [p["text"] for p in kalan] == ["Paragraf 1", "Paragraf 3", "Paragraf 4"]

    # Hepsini sil -> bölüm gerçekten boşalmalı
    for _ in range(3):
        client.delete(f"/chapters/{ch['id']}/paragraphs/1", headers=headers)
    assert client.get(f"/chapters/{ch['id']}", headers=headers).json()["paragraphs"] == []


# --- El yazması dışa aktarımı ----------------------------------------------

def _aktif_kitap(headers):
    """Kitap kimliğini SABİT varsaymak, testler birlikte koşunca kırılıyor -
    fixture'ın verdiği başlıktan oku."""
    return headers["X-Novel-Id"]


def _roman_kur(client, headers):
    """KISIM > bölüm > paragraf yapısı kurar."""
    client.post("/chapters/", json={"number": 1, "title": "BİRİNCİ KISIM", "kind": "part"},
                headers=headers)
    b1 = client.post("/chapters/", json={"number": 2, "title": "Varış"}, headers=headers).json()
    client.put(f"/chapters/{b1['id']}/paragraphs/1",
               json={"number": 1, "text": "İhtiyar teknisyen araçtan indi."}, headers=headers)
    client.put(f"/chapters/{b1['id']}/paragraphs/2",
               json={"number": 2, "text": "Bina güneşte parlıyordu."}, headers=headers)
    b2 = client.post("/chapters/", json={"number": 3, "title": "Bodrum"}, headers=headers).json()
    client.put(f"/chapters/{b2['id']}/paragraphs/1",
               json={"number": 1, "text": "Merdivenler gıcırdadı."}, headers=headers)
    return b1, b2


def test_manuscript_export_markdown_keeps_hierarchy(client, headers):
    """JSON yedeği VERİ yedeğidir - okumak için okunur çıktı gerekiyordu."""
    _roman_kur(client, headers)
    r = client.get(f"/novels/{_aktif_kitap(headers)}/manuscript?format=md", headers=headers)
    assert r.status_code == 200
    metin = r.content.decode("utf-8")
    assert "BİRİNCİ KISIM" in metin and "Varış" in metin
    assert "İhtiyar teknisyen araçtan indi." in metin
    assert "Merdivenler gıcırdadı." in metin
    # Hiyerarşi başlık düzeyine yansımalı: KISIM, bölümden üstte
    assert metin.index("BİRİNCİ KISIM") < metin.index("Varış")
    assert "#" in metin


def test_manuscript_export_docx_is_readable(client, headers):
    from io import BytesIO
    from docx import Document

    _roman_kur(client, headers)
    r = client.get(f"/novels/{_aktif_kitap(headers)}/manuscript?format=docx", headers=headers)
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers["content-type"]
    belge = Document(BytesIO(r.content))
    tumu = "\n".join(p.text for p in belge.paragraphs)
    assert "Bina güneşte parlıyordu." in tumu
    assert any(p.style.name.startswith("Heading") for p in belge.paragraphs)


def test_manuscript_export_txt_and_stats(client, headers):
    _roman_kur(client, headers)
    metin = client.get(f"/novels/{_aktif_kitap(headers)}/manuscript?format=txt", headers=headers).content.decode("utf-8")
    assert "İhtiyar teknisyen araçtan indi." in metin
    assert "#" not in metin, "düz metinde başlık işareti olmamalı"

    ist = client.get(f"/novels/{_aktif_kitap(headers)}/manuscript-stats", headers=headers).json()
    assert ist["paragraf"] == 3
    assert ist["yazili_bolum"] == 2
    assert ist["kelime"] > 5
    assert client.get(f"/novels/{_aktif_kitap(headers)}/manuscript?format=pdf", headers=headers).status_code == 400
