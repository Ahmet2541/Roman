"""YAPI KİLİDİ v1.0 - yapılandırılmış hücre şeması testleri."""
import pytest


@pytest.fixture(autouse=True)
def _reset_rate_limiter():
    from app import ratelimit
    ratelimit._calls.clear()
    yield


def _matris(client, headers):
    return client.post("/matrix/", json={
        "name": "Turlar",
        "columns": [{"label": "TUR 1: BAŞKAN"}],
        "rows": [{"label": "3. Sorgu"}],
    }, headers=headers).json()


def test_structured_cell_renders_to_content(client, headers):
    """data gönderilir, content ondan üretilir - AI'ya giden metin bu."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {
            "olay": "Başkan sanığa mendili uzatır.",
            "zaman": {"tarih": "12 Mart", "saat": "21:40", "tip": "SAYAC"},
            "mekan": "VIP Salonu",
            "kisiler": [{"id": None, "ad": "Vicdan",
                         "duygu": {"baslangic": "güven", "bitis": "şüphe"}}],
            "giris": "İlk repliği reddeder.",
            "gelisme": "Baskı artar.",
            "sonuc": "ÇÖZÜN kelimesi havada kalır.",
            "baglantilar": [{"kod": "mp7", "tur": "ayna", "not": "T1·G1"}],
        },
    }, headers=headers)
    assert r.status_code == 200, r.text
    c = r.json()["content"]
    assert "OLAY: Başkan sanığa mendili uzatır." in c
    assert "ZAMAN: 12 Mart 21:40 (SAYAÇ)" in c
    assert "KİŞİLER: Vicdan (güven → şüphe)" in c
    assert "SONUÇ: ÇÖZÜN kelimesi havada kalır." in c
    assert "BAĞLANTI: MP7 (ayna) → T1·G1" in c


def test_missing_fields_warn_but_do_not_block(client, headers):
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "Bir şey olur."},
    }, headers=headers)
    assert r.status_code == 200, "eksik alan kaydı ENGELLEMEMELİ"
    uyarilar = " · ".join(r.json()["warnings"])
    assert "GİRİŞ" in uyarilar and "MEKAN" in uyarilar


def test_damga_must_appear_in_sonuc(client, headers):
    """Turun damga kelimesi SONUÇ beat'inde geçmiyorsa uyarı çıkar."""
    m = _matris(client, headers)
    col_id = m["columns"][0]["id"]
    client.put(f"/matrix/{m['id']}/columns/{col_id}", json={
        "label": "TUR 1: BAŞKAN", "tur_data": {"damga": "ÇÖZÜN"},
    }, headers=headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "giris": "a", "gelisme": "b", "sonuc": "Kapı kapanır."},
    }, headers=headers)
    assert any("Damga" in w for w in r.json()["warnings"])
    # Damga geçince uyarı kalkar
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "giris": "a", "gelisme": "b", "sonuc": "ÇÖZÜN havada kalır."},
    }, headers=headers)
    assert not any("Damga" in w for w in r2.json()["warnings"])


def test_miras_reaches_ai_context_live(client, headers):
    """TUR/PARÇA mirası hücreye kopyalanmaz - bağlamda canlı okunur."""
    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/generate-chapters", headers=headers)
    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    col_id, row_id = full["columns"][0]["id"], full["rows"][0]["id"]
    cell = full["cells"][0]

    client.put(f"/matrix/{m['id']}/columns/{col_id}", json={
        "label": "TUR 1", "tur_data": {"damga": "ÇÖZÜN", "suc": "ihmal"},
    }, headers=headers)
    client.put(f"/matrix/{m['id']}/rows/{row_id}", json={
        "label": "3. Sorgu", "parca_data": {"no": "3", "sure": "20 dk"},
    }, headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": row_id, "data": {"olay": "Sorgu başlar."},
    }, headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": cell["chapter_number"],
    }, headers=headers).json()["context"]
    assert "OLAY: Sorgu başlar." in ctx
    assert "TUR MİRASI" in ctx and "ÇÖZÜN" in ctx
    assert "PARÇA MİRASI" in ctx and "20 dk" in ctx


def test_chapter_binding_survives_content_only_update(client, headers):
    """chapter_id yollanmayan istek bağı koparmamalı; açık null koparmalı."""
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    m = _matris(client, headers)
    cid, rid = m["columns"][0]["id"], m["rows"][0]["id"]
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": cid, "row_id": rid, "data": {"olay": "ilk"}, "chapter_id": ch["id"],
    }, headers=headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": cid, "row_id": rid, "data": {"olay": "güncel"},
    }, headers=headers)
    assert r.json()["chapter_id"] == ch["id"], "bağ sessizce koptu"
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": cid, "row_id": rid, "data": {"olay": "güncel"}, "chapter_id": None,
    }, headers=headers)
    assert r2.json()["chapter_id"] is None, "açık null bağı koparmalı"


def test_quick_plan_survives_deleted_column(client, headers):
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    client.post("/matrix/quick-plan", json={"chapter_id": ch["id"], "content": "plan"}, headers=headers)
    qm = [x for x in client.get("/matrix/", headers=headers).json() if x["name"] == "Hızlı Planlar"][0]
    full = client.get(f"/matrix/{qm['id']}", headers=headers).json()
    client.delete(f"/matrix/{qm['id']}/columns/{full['columns'][0]['id']}", headers=headers)
    ch2 = client.post("/chapters/", json={"number": 2, "title": "B2"}, headers=headers).json()
    r = client.post("/matrix/quick-plan", json={"chapter_id": ch2["id"], "content": "plan2"}, headers=headers)
    assert r.status_code == 200, r.text


def test_free_text_cells_still_work(client, headers):
    """Eski serbest metin yolu bozulmadı - data göndermeyen istek content yazar."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "content": "serbest plan metni",
    }, headers=headers)
    assert r.json()["content"] == "serbest plan metni"
    assert r.json()["warnings"] == []


# --- v1.1: Qwen denetimi sonrası eklenen alanlar --------------------------

def test_sayac_needs_subject(client, headers):
    """SAYAÇ/ATLAMA tek başına boşlukta durmasın - neyin sayacı yazılmalı."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    r = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "zaman": {"tip": "SAYAC"}},
    }, headers=headers)
    assert any("sayacı" in w for w in r.json()["warnings"])
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "zaman": {"saat": "21:40", "tip": "SAYAC",
                                                 "sayac": "ambulans bekleme süresi"}},
    }, headers=headers)
    assert not any("sayacı" in w for w in r2.json()["warnings"])
    assert "ZAMAN: 21:40 (SAYAÇ: ambulans bekleme süresi)" in r2.json()["content"]


def test_person_without_emotion_is_flagged(client, headers):
    """Yayı olmayan kişi sahnede sadece dekordur."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    r = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "kisiler": [{"ad": "Palyaço"}]},
    }, headers=headers)
    assert any("Palyaço" in w and "duygu" in w for w in r.json()["warnings"])
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "kisiler": [
            {"ad": "Palyaço", "duygu": {"baslangic": "güven", "bitis": "şüphe"}}]},
    }, headers=headers)
    assert "KİŞİLER: Palyaço (güven → şüphe)" in r2.json()["content"]


def test_two_people_two_arcs_in_one_cell(client, headers):
    """Kişi başına yay: iki bilinç sahneyi BÖLMEDEN taşınabilmeli."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "İkisi binaya yürür.", "kisiler": [
            {"ad": "Genç Mühendis", "duygu": {"baslangic": "umut", "bitis": "gurur"}},
            {"ad": "İhtiyar Teknisyen", "duygu": {"baslangic": "nostalji", "bitis": "yorgunluk"}}]},
    }, headers=headers)
    assert ("KİŞİLER: Genç Mühendis (umut → gurur), "
            "İhtiyar Teknisyen (nostalji → yorgunluk)") in r.json()["content"]


def test_multiple_beats_are_numbered(client, headers):
    """Tekli bölümlerde bir aşamada birden çok bağımsız hareket olabilir."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "gelisme": ["İhtiyar mendille terini siler.",
                                          "Genç twit atar."]},
    }, headers=headers)
    metin = r.json()["content"]
    assert "GELİŞME 1: İhtiyar mendille terini siler." in metin
    assert "GELİŞME 2: Genç twit atar." in metin
    # Tek beat varsa numara konmaz - gereksiz gürültü
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "gelisme": ["Tek hareket."]},
    }, headers=headers)
    assert "GELİŞME: Tek hareket." in r2.json()["content"]


def test_parallel_matrix_wants_one_person_one_beat(client, headers):
    """Paralel matriste turlar KARŞILIKLI ilerler: bir turda ihtiyarın
    sorusu varsa ötekinde öğrencininki - tek kişi, tek beat."""
    tekli = client.post("/matrix/", json={
        "name": "Tekli", "columns": [{"label": "K"}], "rows": [{"label": "S"}]},
        headers=headers).json()
    coklu = client.post("/matrix/", json={
        "name": "Paralel", "columns": [{"label": "Tur 1"}, {"label": "Tur 2"}],
        "rows": [{"label": "S"}]}, headers=headers).json()
    veri = {"olay": "x",
            "kisiler": [{"ad": "A", "duygu": {"baslangic": "umut"}},
                        {"ad": "B", "duygu": {"baslangic": "korku"}}],
            "gelisme": ["bir", "iki"]}

    t = client.put(f"/matrix/{tekli['id']}/cells", json={
        "column_id": tekli["columns"][0]["id"], "row_id": tekli["rows"][0]["id"],
        "data": veri}, headers=headers)
    assert not any("Paralel matriste" in w for w in t.json()["warnings"]), \
        "tekli bölümde çoklu yapı serbest olmalı"

    p = client.put(f"/matrix/{coklu['id']}/cells", json={
        "column_id": coklu["columns"][0]["id"], "row_id": coklu["rows"][0]["id"],
        "data": veri}, headers=headers)
    uyarilar = " · ".join(p.json()["warnings"])
    assert "tek kişi olmalı" in uyarilar and "tek beat olmalı" in uyarilar


def test_odak_required_when_multiple_objects(client, headers):
    """Birden çok nesnede dikkat dağılmasın; tek nesnede odak zaten belli."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    coklu = [{"ad": "mendil"}, {"ad": "şişe"}]
    r = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "nesneler": coklu},
    }, headers=headers)
    assert any("ODAK" in w for w in r.json()["warnings"])
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "nesneler": coklu, "odak": "şişe"},
    }, headers=headers)
    assert not any("ODAK" in w for w in r2.json()["warnings"])
    assert "ODAK: şişe" in r2.json()["content"]
    # Tek nesnede uyarı çıkmamalı
    r3 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "nesneler": [{"ad": "mendil"}]},
    }, headers=headers)
    assert not any("ODAK" in w for w in r3.json()["warnings"])


def test_baglanti_must_state_an_action(client, headers):
    """Bağlantı referans etiketi değil, yapılacak iş olmalı."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    r = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "baglantilar": [{"kod": "MP7", "tur": "ayna"}]},
    }, headers=headers)
    assert any("ne yapılacağını" in w for w in r.json()["warnings"])
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "baglantilar": [
            {"kod": "MP7", "tur": "ayna", "not": "MP7'deki ayna imgesini mendille yansıt"}]},
    }, headers=headers)
    assert "MP7 (ayna) → MP7'deki ayna imgesini mendille yansıt" in r2.json()["content"]


# --- v1.2: ikinci Qwen denetimi sonrası ------------------------------------

def test_damga_check_uses_word_boundary(client, headers):
    """Alt dize araması damgayı başka kelimenin içinde bulup uyarıyı
    yutuyordu: "ÇÖZÜN" damgası "çözünürlüğü" kelimesinde eşleşiyor."""
    m = _matris(client, headers)
    col_id, row_id = m["columns"][0]["id"], m["rows"][0]["id"]
    client.put(f"/matrix/{m['id']}/columns/{col_id}", json={
        "label": "TUR 3", "tur_data": {"damga": "ÇÖZÜN"},
    }, headers=headers)
    kandirmaca = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": row_id,
        "data": {"olay": "x", "giris": "a", "gelisme": "b",
                 "sonuc": "raporun çözünürlüğü düşük fotokopisinin hışırtısıdır."},
    }, headers=headers)
    assert any("Damga" in w for w in kandirmaca.json()["warnings"]), \
        "damga başka kelimenin içinde bulunup uyarı yutuldu"
    gercek = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": row_id,
        "data": {"olay": "x", "giris": "a", "gelisme": "b", "sonuc": "ÇÖZÜN masada asılı kalır."},
    }, headers=headers)
    assert not any("Damga" in w for w in gercek.json()["warnings"])


def test_beat_bloat_is_flagged(client, headers):
    """Beat'e olay dizisi yazılırsa yazan modele kuracak yer kalmaz."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    dizi = ("Vicdan üç ayrı soruyla zemin etüdünün neden üç kat geciktiğini sorar; "
            "Jeolog her üçüne de ayrıntılı cevap verir, tarih sıralamasını kendisi "
            "kurar ve raporun kim tarafından imzalandığını açıklar.")
    r = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "gelisme": dizi},
    }, headers=headers)
    assert any("olay dizisi" in w for w in r.json()["warnings"])
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "gelisme": "Jeolog sayfayı çevirmeyi reddeder."},
    }, headers=headers)
    assert not any("olay dizisi" in w for w in r2.json()["warnings"])


# --- Denetim promptu -------------------------------------------------------

def test_audit_prompt_finds_structural_faults(client, headers):
    """Sayılabilir kusurlar deterministik bulunur - modele sorulmaz."""
    m = client.post("/matrix/", json={
        "name": "Turlar", "columns": [{"label": "TUR 1"}],
        "rows": [{"label": "1. Giriş"}, {"label": "2. Sorgu"}],
    }, headers=headers).json()
    col_id = m["columns"][0]["id"]
    # Dolu ama BAĞSIZ plan + var olmayan koda referans
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": m["rows"][0]["id"],
        "data": {"olay": "Bir şey olur.", "baglantilar": [
            {"kod": "MP999", "tur": "ayna", "not": "yansıt"}]},
    }, headers=headers)

    r = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()
    metin = r["prompt"]
    assert "BAĞSIZ PLAN" in metin, "bağlanmamış plan bulunamadı"
    assert "KAYIP REFERANS" in metin and "MP999" in metin
    assert "BOŞ HÜCRE" in metin           # ikinci satır boş
    assert "DAMGASIZ TUR" in metin        # tur_data tanımlanmamış
    assert r["summary"]["dolu_hucre"] == 1
    assert r["summary"]["yapisal_bulgu"] >= 4
    # Sorular da metinde olmalı - kopyalanır kopyalanmaz kullanılabilsin
    assert "SORULAR" in metin and "Çelişki" in metin


def test_audit_prompt_can_scope_to_one_column(client, headers):
    m = client.post("/matrix/", json={
        "name": "T", "columns": [{"label": "TUR 1"}, {"label": "TUR 2"}],
        "rows": [{"label": "1. Giriş"}],
    }, headers=headers).json()
    hedef = m["columns"][0]["id"]
    r = client.get(f"/matrix/{m['id']}/audit-prompt?column_id={hedef}", headers=headers).json()
    assert r["summary"]["kolon_sayisi"] == 1
    assert "TUR 1" in r["prompt"] and "TUR 2" not in r["prompt"]
    assert client.get(f"/matrix/{m['id']}/audit-prompt?column_id=99999",
                      headers=headers).status_code == 404


def test_audit_prompt_clean_matrix_reports_no_faults(client, headers):
    """Temiz planda gürültü üretilmemeli - her plan kusurlu görünmesin."""
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    m = client.post("/matrix/", json={
        "name": "T", "columns": [{"label": "TUR 1"}], "rows": [{"label": "1. Giriş"}],
    }, headers=headers).json()
    col_id = m["columns"][0]["id"]
    client.put(f"/matrix/{m['id']}/columns/{col_id}", json={
        "label": "TUR 1", "tur_data": {"damga": "ÇÖZÜN"},
    }, headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": m["rows"][0]["id"], "chapter_id": ch["id"],
        "data": {"olay": "Sorgu başlar.", "mekan": "Salon",
                 "zaman": {"tip": "NOKTA"},
                 "ortam": {"baslangic": "beklenti"},
                 "kisiler": [{"ad": "Vicdan", "duygu": {"baslangic": "merak"}}],
                 "giris": "Kapı açılır.", "gelisme": "Soru sorulur.",
                 "sonuc": "ÇÖZÜN masada kalır."},
    }, headers=headers)
    r = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()
    assert r["summary"]["yapisal_bulgu"] == 0
    assert "Yapısal kusur bulunamadı" in r["prompt"]
    assert r["summary"]["uyarili_hucre"] == 0


# --- Toplu dışa aktarım ----------------------------------------------------

def _dolu_matris(client, headers):
    ch = client.post("/chapters/", json={"number": 1, "title": "Sorgu"}, headers=headers).json()
    m = client.post("/matrix/", json={
        "name": "Turlar", "columns": [{"label": "TUR 1"}], "rows": [{"label": "1. Giriş"}],
    }, headers=headers).json()
    col_id, row_id = m["columns"][0]["id"], m["rows"][0]["id"]
    client.put(f"/matrix/{m['id']}/columns/{col_id}", json={
        "label": "TUR 1", "tur_data": {"damga": "ÇÖZÜN", "suc": "ihmal"},
    }, headers=headers)
    client.put(f"/matrix/{m['id']}/rows/{row_id}", json={
        "label": "1. Giriş", "instructions": "- Tek cümle",
        "parca_data": {"no": "1", "sure": "20 dk"},
    }, headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": row_id, "chapter_id": ch["id"],
        "data": {"olay": "Sorgu başlar.", "mekan": "Salon"},
    }, headers=headers)
    return m


def test_export_json_carries_everything(client, headers):
    """JSON dökümü yeniden kurulabilecek kadar tam olmalı: miras alanları,
    yapılandırılmış hücre verisi, bölüm bağı, uyarılar."""
    import json as _json
    _dolu_matris(client, headers)
    r = client.get("/matrix/export?format=json", headers=headers)
    assert r.status_code == 200
    assert "attachment" in r.headers["content-disposition"]
    veri = _json.loads(r.content.decode("utf-8"))
    mat = veri["matrisler"][0]
    assert mat["turlar"][0]["tur_mirasi"]["damga"] == "ÇÖZÜN"
    assert mat["asamalar"][0]["parca_mirasi"]["sure"] == "20 dk"
    assert mat["asamalar"][0]["yazim_kisitlari"] == "- Tek cümle"
    h = mat["hucreler"][0]
    assert h["veri"]["olay"] == "Sorgu başlar."
    assert h["bolum"] == 1 and h["kod"]
    assert "OLAY: Sorgu başlar." in h["metin"]
    assert isinstance(h["uyarilar"], list)


def test_export_markdown_is_readable(client, headers):
    _dolu_matris(client, headers)
    r = client.get("/matrix/export?format=md", headers=headers)
    assert r.status_code == 200
    metin = r.content.decode("utf-8")
    assert "# PLAN MATRİSLERİ" in metin
    assert "TUR MİRASI" in metin and "ÇÖZÜN" in metin
    assert "YAZIM KISITLARI: - Tek cümle" in metin
    assert "→ Bölüm 1" in metin


def test_export_scope_and_errors(client, headers):
    m = _dolu_matris(client, headers)
    import json as _json
    tek = _json.loads(client.get(f"/matrix/export?matrix_id={m['id']}",
                                 headers=headers).content.decode("utf-8"))
    assert len(tek["matrisler"]) == 1
    assert client.get("/matrix/export?format=pdf", headers=headers).status_code == 400
    # "export" bir matris kimliği sanılmamalı (yol sırası)
    assert client.get("/matrix/export", headers=headers).status_code == 200


def test_export_keeps_legacy_free_text_cells(client, headers):
    """Eski serbest metinli hücrelerin içeriği dökümde kaybolmamalı."""
    import json as _json
    m = client.post("/matrix/", json={
        "name": "Eski", "columns": [{"label": "K"}], "rows": [{"label": "S"}],
    }, headers=headers).json()
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "content": "serbest plan metni",
    }, headers=headers)
    veri = _json.loads(client.get(f"/matrix/export?matrix_id={m['id']}",
                                  headers=headers).content.decode("utf-8"))
    assert veri["matrisler"][0]["hucreler"][0]["metin"] == "serbest plan metni"


# --- Hedef uzunluk (3 seviye) ---------------------------------------------

def test_uzunluk_three_levels_reach_the_plan_text(client, headers):
    """Etiket değil SOMUT karşılık gitmeli - 'uzun' tek başına modele
    hiçbir şey söylemez."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    beklenen = {
        "ozet": ("ÖZET", "1-2 paragraf"),
        "normal": ("NORMAL", "4-6 paragraf"),
        "uzun": ("UZUN METİN", "8+ paragraf"),
    }
    for seviye, (etiket, olcu) in beklenen.items():
        r = client.put(f"/matrix/{m['id']}/cells", json={
            **ortak, "data": {"olay": "x", "uzunluk": seviye},
        }, headers=headers)
        metin = r.json()["content"]
        assert f"HEDEF UZUNLUK: {etiket}" in metin, f"{seviye} etiketi geçmedi: {metin}"
        assert olcu in metin, f"{seviye} için somut ölçü yok"
        assert r.json()["data"]["uzunluk"] == seviye


def test_uzunluk_defaults_to_normal(client, headers):
    """Belirtilmeyen ya da geçersiz değer normale düşer - uzunluksuz plan,
    modelin her sahnede kendi ölçüsünü seçmesi demek."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    for veri in ({"olay": "x"}, {"olay": "x", "uzunluk": "saçma"}):
        r = client.put(f"/matrix/{m['id']}/cells", json={**ortak, "data": veri}, headers=headers)
        assert r.json()["data"]["uzunluk"] == "normal"
        assert "HEDEF UZUNLUK: NORMAL" in r.json()["content"]


def test_uzunluk_reaches_ai_context(client, headers):
    """Zincirin ucu: seçilen ölçü yazım anında AI'ya gitmeli."""
    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/generate-chapters", headers=headers)
    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    cell = full["cells"][0]
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": cell["column_id"], "row_id": cell["row_id"],
        "data": {"olay": "Sorgu başlar.", "uzunluk": "ozet"},
    }, headers=headers)
    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": cell["chapter_number"],
    }, headers=headers).json()["context"]
    assert "HEDEF UZUNLUK: ÖZET" in ctx and "sahne AÇILMAZ" in ctx


# --- Paralellik denetimi ---------------------------------------------------

def _paralel_matris(client, headers, turlar=3, satirlar=3):
    return client.post("/matrix/", json={
        "name": "Turlar",
        "columns": [{"label": f"Tur {i}"} for i in range(1, turlar + 1)],
        "rows": [{"label": f"{i}. Sahne"} for i in range(1, satirlar + 1)],
    }, headers=headers).json()


def _doldur(client, headers, m, col_idx, row_idx, uzunluk="normal"):
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][col_idx]["id"], "row_id": m["rows"][row_idx]["id"],
        "data": {"olay": "x", "uzunluk": uzunluk},
    }, headers=headers)


def test_untouched_columns_are_not_reported_as_gaps(client, headers):
    """Hiç başlanmamış turu 'eksik' diye raporlamak gerçek delikleri
    gürültüye gömer - 8 turdan 7'si boşken uyarı yağmuru olmamalı."""
    m = _paralel_matris(client, headers)
    for r in range(3):
        _doldur(client, headers, m, 0, r)
    metin = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()["prompt"]
    assert "PARALELLİK DELİĞİ" not in metin, "başlanmamış turlar için uyarı üretildi"


def test_gap_in_started_column_is_reported(client, headers):
    """Asıl korku: Tur 1'de 3 sahne dolu, Tur 2'de 2 - yazımda biri
    eksik sahneyle çıkar."""
    m = _paralel_matris(client, headers)
    for r in range(3):
        _doldur(client, headers, m, 0, r)
    _doldur(client, headers, m, 1, 0)
    _doldur(client, headers, m, 1, 2)   # 2. sahne atlandı
    metin = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()["prompt"]
    delik = [x for x in metin.split("\n") if "PARALELLİK DELİĞİ" in x]
    assert delik and "Tur 2" in delik[0] and "2. Sahne" in delik[0]
    assert not any("Tur 1" in d for d in delik), "eksiksiz tur delik olarak raporlandı"


def test_length_mismatch_across_parallel_cells(client, headers):
    """Aynı satır, farklı ölçü: paralel sahneler farklı boyda çıkar."""
    m = _paralel_matris(client, headers, satirlar=1)
    _doldur(client, headers, m, 0, 0, "ozet")
    _doldur(client, headers, m, 1, 0, "uzun")
    metin = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()["prompt"]
    uyum = [x for x in metin.split("\n") if "UZUNLUK UYUŞMAZLIĞI" in x]
    assert uyum and "1. Sahne" in uyum[0]
    assert "ozet" in uyum[0] and "uzun" in uyum[0]


def test_aligned_columns_report_clean(client, headers):
    m = _paralel_matris(client, headers, satirlar=2)
    for c in (0, 1):
        for r in (0, 1):
            _doldur(client, headers, m, c, r, "normal")
    metin = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()["prompt"]
    assert "PARALELLİK DELİĞİ" not in metin and "UZUNLUK UYUŞMAZLIĞI" not in metin


def test_parallel_findings_also_reach_audit_prompt(client, headers):
    m = _paralel_matris(client, headers, satirlar=2)
    _doldur(client, headers, m, 0, 0)
    _doldur(client, headers, m, 0, 1)
    _doldur(client, headers, m, 1, 0)
    metin = client.get(f"/matrix/{m['id']}/audit-prompt", headers=headers).json()["prompt"]
    assert "PARALELLİK DELİĞİ" in metin


def test_sayac_accepts_turkish_spelling_and_renders_it(client, headers):
    """Kullanıcı 'SAYAÇ' yazsa da anahtar ASCII 'SAYAC' saklanır; ekranda
    ve AI metninde Türkçesi görünür."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "zaman": {"tarih": "12 Mart 2027", "tip": "SAYAÇ",
                                        "sayac": "ambulansın gelişi"}},
    }, headers=headers)
    assert r.json()["data"]["zaman"]["tip"] == "SAYAC"
    assert "ZAMAN: 12 Mart 2027 (SAYAÇ: ambulansın gelişi)" in r.json()["content"]


def test_atlama_warning_asks_the_right_question(client, headers):
    """ATLAMA'da sorulan şey 'neyin sayacı' değil 'neyden atlandığı'."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "zaman": {"tip": "ATLAMA"}},
    }, headers=headers)
    uyari = [w for w in r.json()["warnings"] if "ATLAMA" in w]
    assert uyari and "atlandığı" in uyari[0]


def test_ortam_and_person_emotion_are_separate_arcs(client, headers):
    """Odanın hâli ile kişinin hâli ayrı: asıl değer aradaki farkta.
    Odada gerilim varken Başkan'da korku olması, adamın kalabalığın
    hissettiğinden fazlasını hissettiğini söyler."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "mekan": "VIP Salonu",
                 "ortam": {"baslangic": "endişe", "bitis": "korku"},
                 "kisiler": [{"ad": "Başkan", "duygu": {"baslangic": "soğukkanlılık",
                                                        "bitis": "panik"}}]},
    }, headers=headers)
    metin = r.json()["content"]
    assert "ORTAM: endişe → korku" in metin
    assert "KİŞİLER: Başkan (soğukkanlılık → panik)" in metin
    # ORTAM satırı MEKAN'dan hemen sonra gelmeli - ikisi sahnenin yeri ve havası
    satirlar = metin.split("\n")
    assert satirlar[satirlar.index("MEKAN: VIP Salonu") + 1].startswith("ORTAM:")


def test_identical_ortam_and_person_emotion_is_flagged(client, headers):
    """İki alanın varlık sebebi FARK - aynıysa sahne bunu kullanmıyor."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    ayni = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "ortam": {"baslangic": "gerilim"},
                          "kisiler": [{"ad": "Başkan", "duygu": {"baslangic": "gerilim"}}]},
    }, headers=headers)
    assert any("birebir aynı" in w for w in ayni.json()["warnings"])
    farkli = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "ortam": {"baslangic": "gerilim"},
                          "kisiler": [{"ad": "Başkan", "duygu": {"baslangic": "korku"}}]},
    }, headers=headers)
    assert not any("birebir aynı" in w for w in farkli.json()["warnings"])


def test_empty_ortam_is_warned(client, headers):
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x"},
    }, headers=headers)
    assert any("ORTAM duygusu boş" in w for w in r.json()["warnings"])


# --- Matris etiketi ve sırası -----------------------------------------------

def test_matrix_label_is_derived_from_chapter_bindings(client, headers):
    """Bölüm aralığı ELLE yazılmaz - bağlardan türetilir, böylece bağ
    değişince ad kendiliğinden düzelir."""
    m = client.post("/matrix/", json={
        "name": "Tur Yapısı", "columns": [{"label": "T1"}, {"label": "T2"}],
        "rows": [{"label": "S1"}],
    }, headers=headers).json()
    ozet = lambda: [x for x in client.get("/matrix/", headers=headers).json() if x["id"] == m["id"]][0]
    assert ozet()["chapter_label"] is None, "bağ yokken aralık uydurulmamalı"

    a = client.post("/chapters/", json={"number": 5, "title": "A"}, headers=headers).json()
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "chapter_id": a["id"], "data": {"olay": "x"}}, headers=headers)
    assert ozet()["chapter_label"] == "5. Bölüm", "tek bölümde aralık gösterilmemeli"

    b = client.post("/chapters/", json={"number": 12, "title": "B"}, headers=headers).json()
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][1]["id"], "row_id": m["rows"][0]["id"],
        "chapter_id": b["id"], "data": {"olay": "y"}}, headers=headers)
    o = ozet()
    assert o["chapter_label"] == "5-12. Bölüm"
    assert o["chapter_min"] == 5 and o["chapter_max"] == 12


def test_matrix_order_and_move(client, headers):
    """Araya ekleme: yeni matris sona açılır, sonra yerine taşınır."""
    adlar = ["A", "B", "C"]
    for ad in adlar:
        client.post("/matrix/", json={"name": ad, "columns": [{"label": "K"}],
                                      "rows": [{"label": "S"}]}, headers=headers)
    sirala = lambda: [x["name"] for x in client.get("/matrix/", headers=headers).json()]
    assert sirala() == ["A", "B", "C"]

    c_id = [x for x in client.get("/matrix/", headers=headers).json() if x["name"] == "C"][0]["id"]
    assert client.post(f"/matrix/{c_id}/move?direction=up", headers=headers).json()["moved"]
    assert sirala() == ["A", "C", "B"]
    assert client.post(f"/matrix/{c_id}/move?direction=up", headers=headers).json()["moved"]
    assert sirala() == ["C", "A", "B"]
    # Uçtaki matris daha yukarı gitmez ama hata da vermez
    assert client.post(f"/matrix/{c_id}/move?direction=up", headers=headers).json()["moved"] is False
    assert sirala() == ["C", "A", "B"]
    assert client.post(f"/matrix/{c_id}/move?direction=yan", headers=headers).status_code == 400


def test_plan_cell_entities_reach_the_profile_layer(client, headers):
    """Hücrede yazılı kişinin PROFİLİ de bağlama gitmeli. Eskiden plan
    "KİŞİLER: Genç Mühendis" diyordu ama o kişinin kim olduğu sadece ELLE
    seçilirse gidiyordu - model karakteri tanımadan yazıyordu."""
    k = client.post("/characters/", json={
        "name": "Genç Mühendis", "description": "Yirmi altı yaşında, aceleci."},
        headers=headers).json()
    y = client.post("/places/", json={
        "name": "Lümen Vadisi", "description": "Cam kabuklu araştırma vadisi."},
        headers=headers).json()

    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/generate-chapters", headers=headers)
    cell = client.get(f"/matrix/{m['id']}", headers=headers).json()["cells"][0]
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": cell["column_id"], "row_id": cell["row_id"],
        "data": {"olay": "Varış.", "mekan": "Lümen Vadisi", "mekan_id": y["id"],
                 "kisiler": [{"id": k["id"], "ad": "Genç Mühendis",
                              "duygu": {"baslangic": "umut", "bitis": "gurur"}}]},
    }, headers=headers)

    # selected_entities BOŞ gönderiliyor - profiller yine de gitmeli
    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": cell["chapter_number"],
    }, headers=headers).json()["context"]
    assert "Yirmi altı yaşında, aceleci." in ctx, "kişi profili bağlama gitmedi"
    assert "Cam kabuklu araştırma vadisi." in ctx, "mekan profili bağlama gitmedi"


def test_unregistered_names_do_not_break_context(client, headers):
    """Serbest metin olarak yazılmış (kayıtsız) adlar profil aramaz."""
    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/generate-chapters", headers=headers)
    cell = client.get(f"/matrix/{m['id']}", headers=headers).json()["cells"][0]
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": cell["column_id"], "row_id": cell["row_id"],
        "data": {"olay": "Varış.", "kisiler": [{"ad": "Kayıtsız Kişi"}]},
    }, headers=headers)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": cell["chapter_number"],
    }, headers=headers)
    assert r.status_code == 200
    assert "KİŞİLER: Kayıtsız Kişi" in r.json()["context"]


def test_context_preview_returns_full_prompt(client, headers):
    """Önizleme, Qwen'e giden isteğin TAMAMINI göstermeli: sistem yönergesi
    + bağlam + talimat. Sadece bağlamı göstermek 'AI bunu görüyor mu'
    sorusuna eksik cevap veriyordu."""
    client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1,
        "instruction": "Sahneyi yaz.",
    }, headers=headers).json()
    assert r["system_prompt"], "sistem yönergesi dönmedi"
    assert "=== SİSTEM YÖNERGESİ ===" in r["full_prompt"]
    assert "=== KULLANICI MESAJI ===" in r["full_prompt"]
    assert "CONTEXT:" in r["full_prompt"] and "TALİMAT:" in r["full_prompt"]
    assert "Sahneyi yaz." in r["full_prompt"]
    assert r["context"] in r["full_prompt"], "bağlam tam istekte yok"


def test_preview_matches_what_is_actually_sent(client, headers):
    """Önizleme ile gerçek istek AYNI fonksiyondan kurulmalı - yoksa
    önizleme zamanla gerçekte gidenden sapar ve yalan söyler."""
    from app.qwen_client import build_user_message
    client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1, "instruction": "Yaz.",
    }, headers=headers).json()
    beklenen = build_user_message(r["context"], "Yaz.", None)
    assert beklenen in r["full_prompt"]


def test_context_preview_reports_the_real_cause_on_failure(client, headers, monkeypatch):
    """Bağlam oluşturma patlarsa kullanıcı 'İstek başarısız (500)' değil
    SEBEBİ görmeli - yoksa hangi katmanın bozulduğu ancak sunucu
    günlüğünden anlaşılıyor."""
    from app.routers import ai as ai_router

    def patla(*a, **k):
        raise ValueError("test: fihrist katmanı bozuk")

    monkeypatch.setattr(ai_router, "build_context", patla)
    client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1}, headers=headers)
    assert r.status_code == 500
    detay = r.json()["detail"]
    assert "ValueError" in detay and "fihrist katmanı bozuk" in detay


def test_plan_entity_layer_failure_does_not_kill_context(client, headers, monkeypatch):
    """Plan varlıkları bir İYİLEŞTİRME - patlarsa bağlamın tamamını
    düşürmemeli, elle seçilenlerle devam etmeli."""
    from app import ai_context

    def patla(*a, **k):
        raise RuntimeError("test: plan hücresi okunamadı")

    monkeypatch.setattr(ai_context, "plan_hucre_varliklari", patla)
    client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1}, headers=headers)
    assert r.status_code == 200, "iyileştirme katmanı bütün bağlamı düşürdü"


# --- Fihrist budama ---------------------------------------------------------

BOS_OZET = ("ZAMAN: belirtilmemiş\nOLAY: belirtilmemiş\nMEKAN: belirtilmemiş\n"
            "DEVAMLILIK: Açılış bölümü\nKAPANIŞ TONU: belirtilmemiş")


def _bolum(client, headers, no, baslik, ozet):
    ch = client.post("/chapters/", json={"number": no, "title": baslik}, headers=headers).json()
    client.put(f"/chapters/{ch['id']}", json={"title": baslik, "summary": ozet}, headers=headers)
    return ch


def test_index_drops_empty_summaries(client, headers):
    """Her satırı 'belirtilmemiş' olan özet yer kaplar, bilgi vermez."""
    _bolum(client, headers, 1, "Boş Bölüm", BOS_OZET)
    _bolum(client, headers, 2, "Dolu Bölüm", "OLAY: Adam kapıyı açar ve içeri girer.")
    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 3}, headers=headers).json()["context"]
    # Sadece ÖZET katmanına bak - fihrist HARİTASI'nda başlık görünmeli,
    # orası romanın yapısını gösteriyor, özet değil.
    ozetler = ctx.split("FİHRİST HARİTASI")[0]
    assert "Boş Bölüm" not in ozetler, "içi boş özet fihriste girdi"
    assert "Adam kapıyı açar" in ozetler


def test_index_shortens_distant_chapters(client, headers):
    """Uzak bölümler tek satıra iner, komşular tam kalır."""
    uzun = ("OLAY: Uzaktaki olay tek satıra inmeli.\n"
            "MEKAN: Ayrıntılı mekan tarifi burada uzayıp gider ve çok yer kaplar.\n"
            "ATMOSFER: Bu satır da uzun ve gereksiz yer tutuyor.")
    for no in range(1, 11):
        _bolum(client, headers, no, f"B{no}", uzun.replace("Uzaktaki olay", f"Olay {no}"))

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 9}, headers=headers).json()["context"]
    # Komşu (7-11 arası) tam: ATMOSFER satırı görünür
    assert "Olay 8" in ctx and "Bu satır da uzun" in ctx
    # Uzak (1-6) kısaltılmış: OLAY var, ATMOSFER yok
    fihrist = ctx.split("FİHRİST HARİTASI")[0]
    b1 = [x for x in fihrist.split("\n") if x.startswith("Bölüm 1 ")][0]
    assert "Olay 1" in b1 and "ATMOSFER" not in b1
    assert "uzak bölüm tek satıra indirildi" in ctx


def test_index_trimming_actually_shrinks_context(client, headers):
    """Budama gerçekten yer kazandırmalı - amaç buydu."""
    uzun = "OLAY: " + ("ayrıntı " * 60) + "\nMEKAN: " + ("tarif " * 60)
    for no in range(1, 16):
        _bolum(client, headers, no, f"B{no}", uzun)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 15}, headers=headers).json()
    fihrist = [b for b in r["breakdown"] if "FİHRİST" in b["name"].upper()]
    assert fihrist, "fihrist katmanı yok"
    # 15 bölüm x ~800 karakter budanmadan ~12.000 olurdu
    assert fihrist[0]["char_count"] < 6000, f"budama işe yaramadı: {fihrist[0]['char_count']}"


# --- Fihrist üretimi: alt satırlar sahne, bölüm değil ----------------------

def test_generate_chapters_treats_sub_rows_as_scenes(client, headers):
    """El yazmasındaki yapı: Tur (kısım) → numaralı sahne (bölüm) →
    içindeki beat'ler (sahne). Alt satırlara da bölüm açmak o beat'leri
    fihriste ayrı girdi yapıp bölüm sayısını şişiriyordu."""
    m = client.post("/matrix/", json={
        "name": "Tur Yapısı", "columns": [{"label": "Tur 1"}],
        "rows": [{"label": "1 Hologram", "kind": "main"}],
    }, headers=headers).json()
    for etiket in ("[0] ÇERÇEVE", "GÖRÜNTÜ 1", "GÖRÜNTÜ 2"):
        client.post(f"/matrix/{m['id']}/rows", json={"label": etiket, "kind": "sub"}, headers=headers)
    client.post(f"/matrix/{m['id']}/rows", json={"label": "2 Kamera", "kind": "main"}, headers=headers)

    r = client.post(f"/matrix/{m['id']}/generate-chapters", headers=headers).json()
    assert r["created_parts"] == 1, "kolon bir KISIM olmalı"
    assert r["created_chapters"] == 2, f"sadece ana satırlar bölüm olmalı, {r['created_chapters']} açıldı"

    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    satir_kind = {x["id"]: x["kind"] for x in full["rows"]}
    for c in full["cells"]:
        if satir_kind[c["row_id"]] == "sub":
            assert c["chapter_id"] is None, "alt satır bölüme bağlanmış"
        else:
            assert c["chapter_id"], "ana satır bölüme bağlanmamış"
        assert c["code"], "her hücre MP kodu almalı"


def test_sub_row_plans_reach_their_parent_chapter(client, headers):
    """Bağsız alt sahneler, üstlerindeki bağlı bölümün bağlamına
    'BU BÖLÜMÜN SAHNELERİ' olarak girmeli - zincirin ucu burası."""
    m = client.post("/matrix/", json={
        "name": "T", "columns": [{"label": "Tur 1"}],
        "rows": [{"label": "1 Hologram", "kind": "main"}],
    }, headers=headers).json()
    client.post(f"/matrix/{m['id']}/rows", json={"label": "Alt sahne", "kind": "sub"}, headers=headers)
    client.post(f"/matrix/{m['id']}/generate-chapters", headers=headers)

    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    col_id = full["columns"][0]["id"]
    ana = [r for r in full["rows"] if r["kind"] != "sub"][0]
    alt = [r for r in full["rows"] if r["kind"] == "sub"][0]
    bagli = [c for c in full["cells"] if c["row_id"] == ana["id"]][0]

    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": ana["id"], "data": {"olay": "Ana sahne."}}, headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col_id, "row_id": alt["id"], "data": {"olay": "Alt sahne buraya."}}, headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": bagli["chapter_number"]}, headers=headers).json()["context"]
    assert "Ana sahne." in ctx
    assert "Alt sahne buraya." in ctx, "bağsız alt sahne üst bölümün bağlamına girmedi"


# --- Plan dışına sapma kilidi ----------------------------------------------

def test_odak_is_a_constraint_not_a_hint(client, headers):
    """ODAK 'burayı vurgula' değil 'başkasını yazma' demeli - plan dışı
    nesne ve ayrıntı sızmasının kaynaklarından biri buydu."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "odak": "Lümen binası dış cephesi"},
    }, headers=headers)
    metin = r.json()["content"]
    assert "SADECE bunun üzerinde kalacak" in metin
    assert "başka nesneye, mekana ya da ayrıntıya geçme" in metin


def test_plan_carries_deviation_lock(client, headers):
    """Her plan, sahneye plan dışı kişi/olay sokulmasını yasaklayan
    sınırlarla birlikte gitmeli."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "kisiler": [{"ad": "Genç Mühendis"}]},
    }, headers=headers)
    metin = r.json()["content"]
    assert "SINIRLAR (bu sahne için MUTLAK)" in metin
    assert "Başka" in metin and "karakteri sahneye sokma" in metin
    assert "sahneye TAŞIMA" in metin
    assert "Hedef uzunluğa ulaşmak için olay uydurma" in metin


def test_index_layer_declares_it_is_background_only(client, headers):
    """Fihrist 'malzeme' değil 'geçmiş' olduğunu kendisi söylemeli."""
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    client.put(f"/chapters/{ch['id']}", json={
        "title": "B1", "summary": "OLAY: Vicdan karanlıkta saymaya başladı."}, headers=headers)
    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 2}, headers=headers).json()["context"]
    assert "GEÇMİŞİ ANLAMAN İÇİNDİR" in ctx
    assert "sahneye" in ctx and "TAŞIMA" in ctx


def test_zaman_is_a_constraint_when_clock_given(client, headers):
    """Saat verilmişse günün saatiyle çelişen betimleme yasak - plan
    13:30 derken model 'güneşin son ışınları' yazabiliyordu."""
    m = _matris(client, headers)
    ortak = {"column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"]}
    r = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "zaman": {"tarih": "28 Haziran 2030", "saat": "13:30",
                                                 "tip": "NOKTA"}}}, headers=headers)
    assert "sahne BU AN'da geçer" in r.json()["content"]
    # Saat yoksa uyarı da olmamalı - gereksiz gürültü
    r2 = client.put(f"/matrix/{m['id']}/cells", json={
        **ortak, "data": {"olay": "x", "zaman": {"tarih": "üçüncü gün", "tip": "NOKTA"}}},
        headers=headers)
    assert "sahne BU AN'da geçer" not in r2.json()["content"]


def test_plan_forbids_future_knowledge(client, headers):
    """Vicdan aktif edilmeden 'kayıtları tarıyor' diye yazılıyordu."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x"}}, headers=headers)
    metin = r.json()["content"]
    assert "ZAMAN ÇİZGİSİ" in metin
    assert "Henüz gerçekleşmemiş olayları anlatma" in metin


# --- Yapı analizi sonrası: otorite sırası, gelecek yasağı, fihrist anahtarı ---

def test_plan_is_the_last_layer(client, headers):
    """Plan, modelin en çok uyması gereken katman ama on beşin sekizincisi
    olarak yığının ortasında kalıyordu; fihrist ikinci sıradaydı. Son
    okunan metin en güçlü etkiyi bırakır - plan sona alındı."""
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    client.put(f"/chapters/{ch['id']}", json={
        "title": "B1", "summary": "OLAY: Fihrist satırı."}, headers=headers)
    m = _matris(client, headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "chapter_id": ch["id"], "data": {"olay": "Plan satırı."}}, headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1,
        "include_own_summary": True}, headers=headers).json()["context"]
    assert "BÖLÜM PLANI" in ctx
    # Plan bloğu, fihrist/harita katmanlarının SONRASINDA olmalı
    assert ctx.index("BÖLÜM PLANI") > ctx.index("FİHRİST HARİTASI")


def test_system_prompt_forbids_future_tell_patterns(client, headers):
    """Model geleceği 'henüz bilmiyordu ki' kalıbıyla sızdırıyordu:
    karakterin bilmediğini söyleyerek okura söylüyor. Yasak, kalıbı
    ADIYLA saymalı ve en yüksek otoritede - sistem yönergesinde - olmalı."""
    client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers)
    r = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1}, headers=headers).json()
    sistem = r["system_prompt"]
    assert "PLANA SADAKAT" in sistem and "ZAMAN ÇİZGİSİ" in sistem
    assert "henüz bilmiyordu ki" in sistem
    assert "bir daha asla" in sistem
    assert "Bölüm NUMARASI hikâye sırası DEĞİLDİR" in sistem


def test_index_can_be_switched_off(client, headers):
    """Kronolojik geri sahnelerde fihristteki 'geçmiş' aslında GELECEK -
    tamamen kapatılabilmeli."""
    ch = client.post("/chapters/", json={"number": 1, "title": "B1"}, headers=headers).json()
    client.put(f"/chapters/{ch['id']}", json={
        "title": "B1", "summary": "OLAY: Vicdan karanlıkta saymaya başladı."}, headers=headers)
    client.post("/chapters/", json={"number": 2, "title": "B2"}, headers=headers)

    acik = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 2}, headers=headers).json()["context"]
    assert "Vicdan karanlıkta" in acik

    kapali = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 2,
        "include_index": False}, headers=headers).json()["context"]
    assert "Vicdan karanlıkta" not in kapali, "fihrist kapatılmadı"
    assert "ROMAN FİHRİSTİ" not in kapali
    # Diğer katmanlar yerinde kalmalı - sadece fihrist kapanır
    assert "FİHRİST HARİTASI" in kapali


# --- Kronolojik süzme -------------------------------------------------------

def test_chapters_set_later_in_story_are_excluded(client, headers):
    """Bölüm NUMARASI hikâye sırası değil. 28 Haziran'da geçen sahne
    Bölüm 3'teyse, Bölüm 1-2'de anlatılan 7 Temmuz olayları o sahnenin
    GELECEĞİDİR - fihristten çıkmalı."""
    gec = client.post("/chapters/", json={"number": 1, "title": "Yargı"}, headers=headers).json()
    client.put(f"/chapters/{gec['id']}", json={
        "title": "Yargı", "summary": "OLAY: Vicdan sekiz kişiyi masaya oturttu."}, headers=headers)
    erken = client.post("/chapters/", json={"number": 2, "title": "Varış"}, headers=headers).json()

    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/rows", json={"label": "S2"}, headers=headers)
    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    col = full["columns"][0]["id"]
    # Bölüm 1 -> 7 Temmuz (geç), Bölüm 2 -> 28 Haziran (erken)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col, "row_id": full["rows"][0]["id"], "chapter_id": gec["id"],
        "data": {"olay": "x", "zaman": {"tarih": "7 Temmuz 2030", "saat": "21:00", "tip": "NOKTA"}}},
        headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col, "row_id": full["rows"][1]["id"], "chapter_id": erken["id"],
        "data": {"olay": "y", "zaman": {"tarih": "28 Haziran 2030", "saat": "13:30", "tip": "NOKTA"}}},
        headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 2}, headers=headers).json()["context"]
    ozetler = ctx.split("FİHRİST HARİTASI")[0]
    assert "Vicdan sekiz kişiyi" not in ozetler, "gelecekte geçen bölüm fihriste girdi"
    assert "bu sahneden SONRA geçtiği için" in ctx


def test_earlier_chapters_still_reach_context(client, headers):
    """Süzme sadece GELECEĞİ keser - geçmiş yerinde kalmalı."""
    erken = client.post("/chapters/", json={"number": 1, "title": "Varış"}, headers=headers).json()
    client.put(f"/chapters/{erken['id']}", json={
        "title": "Varış", "summary": "OLAY: İkisi binaya girdi."}, headers=headers)
    gec = client.post("/chapters/", json={"number": 2, "title": "Yargı"}, headers=headers).json()

    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/rows", json={"label": "S2"}, headers=headers)
    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    col = full["columns"][0]["id"]
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col, "row_id": full["rows"][0]["id"], "chapter_id": erken["id"],
        "data": {"olay": "x", "zaman": {"tarih": "28 Haziran 2030", "saat": "13:30", "tip": "NOKTA"}}},
        headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col, "row_id": full["rows"][1]["id"], "chapter_id": gec["id"],
        "data": {"olay": "y", "zaman": {"tarih": "7 Temmuz 2030", "saat": "21:00", "tip": "NOKTA"}}},
        headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 2}, headers=headers).json()["context"]
    assert "İkisi binaya girdi." in ctx, "geçmiş bölüm de süzüldü"


def test_undated_chapters_are_not_filtered(client, headers):
    """Zamanı çözülemeyen bölüm süzmenin DIŞINDA kalır - uydurma tarihle
    yanlış sıralamak, sıralamamaktan kötüdür."""
    tarihsiz = client.post("/chapters/", json={"number": 1, "title": "Tarihsiz"}, headers=headers).json()
    client.put(f"/chapters/{tarihsiz['id']}", json={
        "title": "Tarihsiz", "summary": "OLAY: Zamanı belirsiz bir sahne."}, headers=headers)
    simdi = client.post("/chapters/", json={"number": 2, "title": "Şimdi"}, headers=headers).json()

    m = _matris(client, headers)
    client.post(f"/matrix/{m['id']}/rows", json={"label": "S2"}, headers=headers)
    full = client.get(f"/matrix/{m['id']}", headers=headers).json()
    col = full["columns"][0]["id"]
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col, "row_id": full["rows"][0]["id"], "chapter_id": tarihsiz["id"],
        "data": {"olay": "x", "zaman": {"tarih": "üçüncü gün", "tip": "NOKTA"}}}, headers=headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": col, "row_id": full["rows"][1]["id"], "chapter_id": simdi["id"],
        "data": {"olay": "y", "zaman": {"tarih": "28 Haziran 2030", "tip": "NOKTA"}}}, headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 2}, headers=headers).json()["context"]
    assert "Zamanı belirsiz bir sahne." in ctx


def test_export_docx_is_a_real_word_file(client, headers):
    """Word dökümü gerçek bir .docx olmalı ve plan içeriğini taşımalı."""
    from io import BytesIO
    from docx import Document

    _dolu_matris(client, headers)
    r = client.get("/matrix/export?format=docx", headers=headers)
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers["content-type"]
    assert ".docx" in r.headers["content-disposition"]

    belge = Document(BytesIO(r.content))
    metinler = [p.text for p in belge.paragraphs]
    tumu = "\n".join(metinler)
    assert "Plan Matrisleri" in tumu
    assert "TUR 1" in tumu and "1. Giriş" in tumu
    assert "ÇÖZÜN" in tumu                    # tur mirası
    assert "20 dk" in tumu                    # parça mirası
    assert "YAZIM KISITLARI" in tumu
    assert "Sorgu başlar." in tumu            # hücre planı
    assert "Bölüm 1" in tumu                  # bölüm bağı

    # Başlıklar GERÇEK Word başlığı olmalı - belge içi gezinme çalışsın
    basliklar = [p.style.name for p in belge.paragraphs if p.style.name.startswith("Heading")]
    assert basliklar, "hiç Word başlığı yok"


def test_export_docx_marks_empty_cells(client, headers):
    """Boş kesişim basılı planda da görünmeli."""
    from io import BytesIO
    from docx import Document

    m = client.post("/matrix/", json={
        "name": "T", "columns": [{"label": "Tur 1"}],
        "rows": [{"label": "S1"}, {"label": "S2"}]}, headers=headers).json()
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "Dolu sahne."}}, headers=headers)

    r = client.get(f"/matrix/export?format=docx&matrix_id={m['id']}", headers=headers)
    tumu = "\n".join(p.text for p in Document(BytesIO(r.content)).paragraphs)
    assert "Dolu sahne." in tumu
    assert "(boş)" in tumu


def test_export_rejects_unknown_format(client, headers):
    _dolu_matris(client, headers)
    r = client.get("/matrix/export?format=pdf", headers=headers)
    assert r.status_code == 400
    assert "docx" in r.json()["detail"]


def test_human_exports_strip_ai_directives(client, headers):
    """content MODELE giden metindir: SINIRLAR bloğu, 'sahne BU AN'da
    geçer' uyarısı, hedef uzunluğun tarifi hep modele söylenmiş şeyler.
    Basılı planda bunlar gürültü - yazar kendi planını okumak ister."""
    from io import BytesIO
    from docx import Document

    m = _matris(client, headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "Sorgu.", "odak": "mendil",
                 "zaman": {"tarih": "28 Haziran 2030", "saat": "13:30", "tip": "NOKTA"},
                 "nesneler": [{"ad": "mendil"}, {"ad": "şişe"}]},
    }, headers=headers)

    md = client.get("/matrix/export?format=md", headers=headers).content.decode("utf-8")
    dx = "\n".join(p.text for p in Document(
        BytesIO(client.get("/matrix/export?format=docx", headers=headers).content)).paragraphs)

    for cikti, ad in ((md, "markdown"), (dx, "docx")):
        assert "SINIRLAR" not in cikti, f"{ad}: AI yönergesi sızdı"
        assert "sahne BU AN'da geçer" not in cikti, f"{ad}: zaman uyarısı sızdı"
        assert "betimleme SADECE" not in cikti, f"{ad}: odak uyarısı sızdı"
        # Bilgi KALMALI - sadece yönerge kuyruğu kesiliyor
        assert "28 Haziran 2030 13:30" in cikti, f"{ad}: zaman bilgisi kayboldu"
        assert "ODAK: mendil" in cikti, f"{ad}: odak kayboldu"
        assert "HEDEF UZUNLUK: NORMAL" in cikti, f"{ad}: uzunluk etiketi kayboldu"


def test_ai_context_still_has_the_directives(client, headers):
    """Temizlik SADECE okunur çıktılarda - modele giden metin dokunulmaz."""
    m = _matris(client, headers)
    r = client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "odak": "mendil"}}, headers=headers)
    assert "SINIRLAR" in r.json()["content"]
    assert "betimleme SADECE" in r.json()["content"]


# --- Varoluş aralığı: varlık bu sahnede var mı? ---------------------------

def _zamanli_bolum(client, headers, tarih, saat=""):
    """Belirli bir hikâye zamanına bağlı bölüm kurar."""
    ch = client.post("/chapters/", json={"number": 1, "title": "Sahne"}, headers=headers).json()
    m = _matris(client, headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "chapter_id": ch["id"],
        "data": {"olay": "x", "zaman": {"tarih": tarih, "saat": saat, "tip": "NOKTA"}},
    }, headers=headers)
    return ch


def test_entity_not_yet_existing_is_flagged(client, headers):
    """'Vicdan aktifleşmeden kayıtları tarıyor' bir ÜSLUP hatası değil
    VARLIK hatasıydı - o anda henüz yoktu."""
    v = client.post("/characters/", json={
        "name": "Vicdan", "description": "Yapay zekâ.",
        "var_olus": "28 Haziran 2030"}, headers=headers).json()
    _zamanli_bolum(client, headers, "20 Haziran 2030")

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [{"entity_type": "character", "entity_id": v["id"]}],
        "chapter_number": 1}, headers=headers).json()["context"]
    assert "HENÜZ YOK" in ctx
    assert "adı anılamaz" in ctx


def test_entity_already_gone_is_flagged(client, headers):
    k = client.post("/characters/", json={
        "name": "Leyla", "description": "Eşi.",
        "yok_olus": "1 Ocak 2027"}, headers=headers).json()
    _zamanli_bolum(client, headers, "28 Haziran 2030")

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [{"entity_type": "character", "entity_id": k["id"]}],
        "chapter_number": 1}, headers=headers).json()["context"]
    assert "ARTIK YOK" in ctx
    assert "yalnızca hatırlanabilir" in ctx


def test_entity_within_lifespan_is_not_flagged(client, headers):
    """Aralığın içindeyse uyarı YOK - sadece bilgi olarak yazılır."""
    k = client.post("/characters/", json={
        "name": "Genç Mühendis", "description": "Mühendis.",
        "var_olus": "3 Mayıs 2004"}, headers=headers).json()
    _zamanli_bolum(client, headers, "28 Haziran 2030")

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [{"entity_type": "character", "entity_id": k["id"]}],
        "chapter_number": 1}, headers=headers).json()["context"]
    assert "HENÜZ YOK" not in ctx and "ARTIK YOK" not in ctx
    assert "Varoluş: 3 Mayıs 2004" in ctx


def test_unparseable_lifespan_is_left_alone(client, headers):
    """'yedi yıl önce' çözülemez - uydurma tarihle yanlış hüküm vermektense
    denetimin dışında kalmalı."""
    k = client.post("/characters/", json={
        "name": "Belirsiz", "description": "x", "var_olus": "yedi yıl önce"},
        headers=headers).json()
    _zamanli_bolum(client, headers, "28 Haziran 2030")

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [{"entity_type": "character", "entity_id": k["id"]}],
        "chapter_number": 1}, headers=headers).json()["context"]
    assert "HENÜZ YOK" not in ctx


def test_place_and_object_have_lifespan_too(client, headers):
    """Lümen binası yedi yıl önce yıkılan binanın yerine dikildi -
    enkaz sahnelerinde yeni bina YOK."""
    y = client.post("/places/", json={
        "name": "Lümen Binası", "description": "Cam kabuk.",
        "var_olus": "1 Haziran 2030"}, headers=headers).json()
    n = client.post("/objects/", json={
        "name": "Süper bilgisayar", "description": "Kabin.",
        "var_olus": "28 Haziran 2030"}, headers=headers).json()
    _zamanli_bolum(client, headers, "1 Ocak 2029")

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [{"entity_type": "place", "entity_id": y["id"]},
                              {"entity_type": "object", "entity_id": n["id"]}],
        "chapter_number": 1}, headers=headers).json()["context"]
    assert ctx.count("HENÜZ YOK") == 2


def test_lifespan_not_added_to_factions_or_terms(client, headers):
    """Varoluş aralığı SADECE kişi/mekan/nesne için - gruplar ve terimler
    zamansız kavramlar."""
    r = client.post("/factions/", json={"name": "Grup", "var_olus": "2030"}, headers=headers)
    assert r.status_code in (200, 201)
    assert "var_olus" not in r.json()


def test_aliases_reach_the_context(client, headers):
    """Takma adlar yalnızca anma tespitinde kullanılıyordu. Model
    "Vicdan'a Ayna da deniyor" bilgisini bilmezse o adı kullanamaz ve
    metinde geçtiğinde aynı kişi olduğunu anlamaz."""
    k = client.post("/characters/", json={
        "name": "İhtiyar Teknisyen", "description": "Elektronikçi.",
        "aliases": ["usta", "ihtiyar"]}, headers=headers).json()
    client.post("/chapters/", json={"number": 1, "title": "S"}, headers=headers)
    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [{"entity_type": "character", "entity_id": k["id"]}],
        "chapter_number": 1}, headers=headers).json()["context"]
    assert "Ayrıca şöyle anılır: usta, ihtiyar" in ctx


def test_meta_never_reaches_context_even_in_hidden_mode(client, headers):
    """Meta SADECE yazar içindir - alt-metin modu açıkken bile gitmemeli."""
    k = client.post("/characters/", json={
        "name": "Vicdan", "description": "x",
        "sections": {"meta": "Sembol: adalet terazisi.", "gizli": "İntikam peşinde."}},
        headers=headers).json()
    client.post("/chapters/", json={"number": 1, "title": "S"}, headers=headers)
    ref = [{"entity_type": "character", "entity_id": k["id"]}]
    for gizli in (False, True):
        ctx = client.post("/ai/context-preview", json={
            "selected_entities": ref, "chapter_number": 1,
            "include_hidden": gizli}, headers=headers).json()["context"]
        assert "adalet terazisi" not in ctx, f"meta sızdı (include_hidden={gizli})"
    # Gizli katman ise SADECE alt-metin modunda gitmeli
    kapali = client.post("/ai/context-preview", json={
        "selected_entities": ref, "chapter_number": 1}, headers=headers).json()["context"]
    acik = client.post("/ai/context-preview", json={
        "selected_entities": ref, "chapter_number": 1,
        "include_hidden": True}, headers=headers).json()["context"]
    assert "İntikam peşinde." not in kapali
    assert "İntikam peşinde." in acik


# --- İçe aktarım -----------------------------------------------------------

def test_import_roundtrip_recreates_the_matrix(client, headers):
    """Dışa aktar → içe aktar: yapı ve içerik korunmalı."""
    import json as _json
    _dolu_matris(client, headers)
    veri = _json.loads(client.get("/matrix/export?format=json",
                                  headers=headers).content.decode("utf-8"))

    r = client.post("/matrix/import", json=veri, headers=headers).json()
    assert r["matris"] == 1 and r["kolon"] == 1 and r["satir"] == 1 and r["hucre"] == 1

    liste = client.get("/matrix/", headers=headers).json()
    yeni = [m for m in liste if "(içe aktarıldı)" in m["name"]]
    assert yeni, "aynı adlı matris ayırt edilmedi"
    tam = client.get(f"/matrix/{yeni[0]['id']}", headers=headers).json()
    assert tam["columns"][0]["tur_data"]["damga"] == "ÇÖZÜN"
    assert tam["rows"][0]["parca_data"]["sure"] == "20 dk"
    assert tam["rows"][0]["instructions"] == "- Tek cümle"
    assert "OLAY: Sorgu başlar." in tam["cells"][0]["content"]


def test_import_never_overwrites_existing(client, headers):
    """Mevcut matrisin ÜZERİNE YAZMAMALI - yeni matris eklenmeli."""
    import json as _json
    m = _dolu_matris(client, headers)
    veri = _json.loads(client.get("/matrix/export?format=json",
                                  headers=headers).content.decode("utf-8"))
    client.post("/matrix/import", json=veri, headers=headers)

    eski = client.get(f"/matrix/{m['id']}", headers=headers).json()
    assert "OLAY: Sorgu başlar." in eski["cells"][0]["content"], "orijinal bozuldu"
    assert len(client.get("/matrix/", headers=headers).json()) == 2


def test_import_reresolves_entities_by_name(client, headers):
    """Dosyadaki ID'ler KAYNAK evrene ait - körlemesine kullanmak yanlış
    kayda bağlar. Hedef evrende ADLA yeniden çözülmeli."""
    import json as _json
    k = client.post("/characters/", json={"name": "Genç Mühendis"}, headers=headers).json()
    m = _matris(client, headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "data": {"olay": "x", "kisiler": [
            {"id": 9999, "ad": "Genç Mühendis"},          # YANLIŞ id
            {"id": 8888, "ad": "Kayıtsız Kişi"}]}}, headers=headers)
    veri = _json.loads(client.get("/matrix/export?format=json",
                                  headers=headers).content.decode("utf-8"))

    r = client.post("/matrix/import", json=veri, headers=headers).json()
    assert r["cozulen"] >= 1 and r["cozulemeyen"] >= 1

    yeni = [x for x in client.get("/matrix/", headers=headers).json()
            if "(içe aktarıldı)" in x["name"]][0]
    hucre = client.get(f"/matrix/{yeni['id']}", headers=headers).json()["cells"][0]
    kisiler = {p["ad"]: p["id"] for p in hucre["data"]["kisiler"]}
    assert kisiler["Genç Mühendis"] == k["id"], "ID adla yeniden çözülmedi"
    assert kisiler["Kayıtsız Kişi"] is None, "bulunamayan ad için ölü ID kaldı"


def test_import_leaves_missing_chapters_unbound(client, headers):
    """O numarada bölüm yoksa sessizce yanlış bölüme bağlamaktansa
    bağsız bırakmalı - ve bunu raporlamalı."""
    import json as _json
    _dolu_matris(client, headers)
    veri = _json.loads(client.get("/matrix/export?format=json",
                                  headers=headers).content.decode("utf-8"))
    veri["matrisler"][0]["hucreler"][0]["bolum"] = 999

    r = client.post("/matrix/import", json=veri, headers=headers).json()
    assert r["baglanamayan"] == 1
    assert any("999" in u for u in r["uyarilar"])


def test_import_rejects_wrong_file(client, headers):
    r = client.post("/matrix/import", json={"foo": "bar"}, headers=headers)
    assert r.status_code == 400
    assert "matrisler" in r.json()["detail"]


def test_plan_characters_always_bring_their_voice(client, headers):
    """Bölüm seçimi talimattaki anahtar kelimeye bakıyor; "bu sahneyi yaz"
    talimatında "konuş/replik" geçmez. Ama PLANDA yazılı kişiler o sahnede
    KONUŞACAK kişilerdir - sesleri gitmezse model hepsini aynı ağızdan yazar."""
    k = client.post("/characters/", json={
        "name": "İhtiyar Teknisyen", "description": "Elektronikçi.",
        "sections": {"konusma_tarzi": "Soruya soruyla karşılık verir.",
                     "duygusal_yapi": "Suçluluk taşır ama dile getirmez.",
                     "fiziksel_yapi": "Gümüş sakallı."}}, headers=headers).json()
    ch = client.post("/chapters/", json={"number": 1, "title": "S"}, headers=headers).json()
    m = _matris(client, headers)
    client.put(f"/matrix/{m['id']}/cells", json={
        "column_id": m["columns"][0]["id"], "row_id": m["rows"][0]["id"],
        "chapter_id": ch["id"],
        "data": {"olay": "x", "kisiler": [
            {"id": k["id"], "ad": "İhtiyar Teknisyen",
             "duygu": {"baslangic": "umut", "bitis": "keder"}}]}}, headers=headers)

    ctx = client.post("/ai/context-preview", json={
        "selected_entities": [], "chapter_number": 1,
        "instruction": "Bu sahneyi yaz."}, headers=headers).json()["context"]
    assert "Soruya soruyla karşılık verir." in ctx, "plandaki kişinin sesi gitmedi"
    assert "Suçluluk taşır" in ctx, "iç dünyası gitmedi"
    # Seçicilik korunmalı: ilgisiz bölüm YİNE de sadece isim olarak listelenir
    assert "Gümüş sakallı." not in ctx, "seçicilik bozuldu, her şey gidiyor"
